"""Generate the GiveWiZe Technical Report rough draft as a Word document."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)


def add_bullets(items, style_name="List Bullet"):
    for item in items:
        doc.add_paragraph(item, style=style_name)


def add_numbered(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


# ── Title ────────────────────────────────────────────────────────────

title = doc.add_heading(
    "GiveWiZe: An AI-Augmented Charity Discovery and Digital Marketing Platform",
    level=0,
)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Abstract ─────────────────────────────────────────────────────────

doc.add_heading("Abstract", level=1)
doc.add_paragraph(
    "The charitable giving landscape presents significant challenges for donors: "
    "over 1.5 million registered 501(c)(3) nonprofits in the US create information "
    "overload, a persistent trust gap leaves donors uncertain about where their money "
    "goes, and existing platforms offer no personalized matching. GiveWiZe addresses "
    "these barriers through an AI-powered charity discovery platform that aggregates "
    "data from multiple sources, computes a transparent weighted scoring system, "
    "matches donors to charities via an intelligent quiz, and automates charity "
    "onboarding through a pipeline of three Claude AI agents. Additionally, a "
    "Python-based social media content generator produces branded marketing content "
    "and auto-posts to LinkedIn, bridging the gap between charity data and donor "
    "engagement. The platform currently hosts 76+ verified charities across 18 cause "
    "categories, with the automated pipeline processing a new charity in approximately "
    "15-20 seconds at a cost of $0.05 per charity."
)

doc.add_page_break()

# ── 1. Introduction / Problem Definition ─────────────────────────────

doc.add_heading("1. Introduction / Problem Definition", level=1)

doc.add_heading("1.1 The Problem", level=2)
doc.add_paragraph(
    "There are over 1.5 million registered 501(c)(3) nonprofits in the United States. "
    "For donors who want to make a meaningful impact with their charitable giving, "
    "this abundance creates three distinct barriers:"
)
add_bullets([
    "Information Overload: With millions of options and no centralized discovery tool, "
    "donors are overwhelmed before they even begin researching. Most give to charities "
    "they already know rather than finding the best fit for their values.",
    "Trust Gap: According to BBB surveys, only 10% of Americans trust charities "
    '"a lot." Donors don\'t know where their money goes, how efficiently it\'s used, '
    "or whether an organization is legitimate.",
    "No Personalization: Existing platforms like Charity Navigator and GuideStar rate "
    "charities on financial metrics but offer no way to match a charity to an individual "
    "donor's values, interests, or giving preferences.",
])
doc.add_paragraph(
    "An additional problem exists on the marketing side: most small and mid-size "
    "nonprofits lack the resources for consistent social media presence and donor "
    "outreach, while platforms that serve donors have no built-in tools for "
    "content-driven engagement."
)

doc.add_heading("1.2 Existing Solutions and Their Gaps", level=2)
add_bullets([
    "Charity Navigator: Rates financial health and accountability but offers no "
    "personalization or matching. Users must already know what they are looking for.",
    "GuideStar/Candid: Provides raw nonprofit data as a data warehouse, but is not "
    "consumer-friendly and requires significant expertise to interpret.",
    "GoFundMe/JustGiving: Fundraising platforms focused on individual campaigns, not "
    "charity discovery or vetting.",
])
doc.add_paragraph(
    "No existing solution combines transparent scoring, personalized matching, "
    "automated onboarding, and AI-driven marketing content generation in a single platform."
)

doc.add_heading("1.3 Proposed Solution", level=2)
doc.add_paragraph("GiveWiZe is an all-in-one charity discovery and marketing platform that:")
add_numbered([
    "Aggregates data from multiple sources (ProPublica, Charity Navigator, GuideStar, "
    "BBB, charity websites)",
    "Computes a transparent, weighted GiveWiZe Score for every charity",
    "Matches donors to charities via an AI-powered progressive quiz",
    "Auto-onboards new charities through a fully automated AI pipeline with three Claude agents",
    "Generates branded social media content with a Python-based AI marketing tool",
    "Auto-posts to LinkedIn (and Instagram when available)",
])

doc.add_heading("1.4 Target Audience", level=2)
add_bullets([
    "Individual donors who want to give but don't know where to start",
    "Students and young professionals exploring philanthropy for the first time",
    "Nonprofit professionals interested in transparency tools and visibility",
])

doc.add_page_break()

# ── 2. Technical Implementation ──────────────────────────────────────

doc.add_heading("2. Technical Implementation", level=1)

doc.add_heading("2.1 System Architecture Overview", level=2)
doc.add_paragraph("The GiveWiZe platform consists of four major layers:")
add_bullets([
    "Frontend: React 18 + TypeScript + Vite + Tailwind CSS + shadcn/ui components",
    "Backend: Supabase (PostgreSQL database, Authentication, Edge Functions, Storage)",
    "AI Layer: Claude Sonnet (3 agents in Edge Functions) + Claude for content generation",
    "Marketing Tool: Python script with Anthropic API, DALL-E image generation, LinkedIn API",
    "Email: Resend API for transactional emails with custom domain (givewize.org)",
    "Data Sources: ProPublica Nonprofit Explorer API, Charity Navigator (web scraper), charity websites",
])
doc.add_paragraph("[INSERT: System architecture diagram showing data flow between all components]")

doc.add_heading("2.2 Database Design", level=2)
doc.add_paragraph("The application uses Supabase PostgreSQL with 6 tables:")
add_bullets([
    "charities (76 rows) - core charity profiles with 30+ fields including scores, programs, financial data",
    "charity_requests - onboarding pipeline tracking with pipeline_metrics JSONB column",
    "online_review_sources (4 per charity) - ratings from Charity Navigator, GuideStar, BBB, GreatNonprofits",
    "community_reviews - user-submitted reviews with approval workflow",
    "user_profiles - donor accounts with bio, avatar, preferences",
    "user_favorites - saved charities per user",
])
doc.add_paragraph(
    "Row Level Security (RLS) enforces access control at the database level. The frontend "
    "uses an anonymous key (RLS-enforced) while Edge Functions use a service role key that "
    "bypasses RLS for administrative operations."
)

doc.add_heading("2.3 Data Processing Pipeline", level=2)
doc.add_paragraph(
    "ProPublica API: Verifies 501(c)(3) status by EIN, retrieves tax filing data "
    "including revenue, assets, and NTEE classification codes."
)
doc.add_paragraph(
    "Charity Navigator Scraper: Fetches charitynavigator.org/ein/{ein}, extracts star "
    "rating (0-4 scale) and raw expense dollar amounts from embedded __NEXT_DATA__ JSON. "
    "Computes administrative and fundraising expense percentages from raw dollar amounts "
    "since the site does not render these as text."
)
doc.add_paragraph(
    "Website Scraper: Fetches up to 7 pages per charity (homepage, /about, /contact, "
    "/programs, /what-we-do, /services, /our-work) in parallel. Extracts meta descriptions, "
    "campaign keywords, contact information, and donation CTAs."
)

doc.add_heading("2.4 Scoring Methodology", level=2)
doc.add_paragraph("The GiveWiZe Score (0-5 scale) consists of four weighted sub-scores:")
add_bullets([
    "Financial Efficiency (35%) - based on program expense percentage brackets",
    "Transparency (30%) - points awarded for website, published financials, community rating, and EIN verification",
    "Impact (25%) - coverage ratio: people_served_annually / target_population_size",
    "Longevity (10%) - years since founding, with a floor at 3.0 so newer organizations are not heavily penalized",
])
doc.add_paragraph(
    "A critical design decision: null sub-scores are excluded and remaining weights are "
    "renormalized rather than treating missing data as zero. This prevents charities from "
    "being unfairly penalized for incomplete data."
)
doc.add_paragraph(
    "The Community Rating (0-5 scale) aggregates five factors: Online Reputation (30%), "
    "Donor Satisfaction (25%), Transparency (20%), Impact (15%), and Communication (10%). "
    "The Combined Score is calculated as (GiveWiZe Score x 0.6) + (Community Rating x 0.4). "
    "All scores are computed client-side in real-time."
)
doc.add_paragraph("[INSERT: Scoring formula diagram with weights and sub-scores]")

doc.add_heading("2.5 AI Quiz Matching Algorithm", level=2)
doc.add_paragraph("The quiz uses three progressive tiers:")
add_bullets([
    "Quick Match (3 questions) - broad matching in under 60 seconds",
    "Refined Match (6 questions) - deeper value alignment",
    "Deep Match (8 questions) - comprehensive personality-to-charity matching",
])
doc.add_paragraph(
    "The algorithm evaluates 8 matching dimensions: cause area, geographic preference, "
    "organization size, financial priority, impact type, donor experience, engagement style, "
    "and time horizon. Each dimension is weighted (cause area: 5, donor experience: 2, etc.) "
    "and produces a 0-1 match score. Results include ranked charities with 0-100% match "
    "percentages and natural language explanations of why each charity matches."
)

doc.add_heading("2.6 AI-Powered Charity Onboarding Pipeline", level=2)
doc.add_paragraph(
    "The automated onboarding pipeline processes a new charity in approximately 15-20 "
    "seconds through 10 steps:"
)
add_numbered([
    "Duplicate check (name + EIN)",
    "ProPublica EIN verification",
    "Data mapping (ProPublica fields to charity schema)",
    "Parallel enrichment: Charity Navigator scrape + Website scrape",
    "Agent 1 (Enrichment): Claude Sonnet extracts mission, programs, logo, contact info from raw HTML",
    "Data merge (database fields + scraped data + AI extraction)",
    "Score computation",
    "Agent 2 (Quality Review): Claude reviews assembled profile, flags issues, can override category",
    "Auto-approve or flag for manual review",
    "Insert charity + 4 review sources + Agent 3 (Email): personalized outreach for missing data",
])
doc.add_paragraph(
    "Each AI agent has a non-AI fallback ensuring 100% pipeline completion. Agent 1 falls "
    "back to regex extraction, Agent 2 uses default approval logic, and Agent 3 uses a "
    "template email. The pipeline costs approximately $0.05 per charity for three Claude "
    "API calls. Flagged charities trigger an admin notification email to info@givewize.org."
)
doc.add_paragraph("[INSERT: Onboarding pipeline flowchart with AI agent callouts]")
doc.add_paragraph("[INSERT: Sample production AI agent prompt]")

doc.add_heading("2.7 AI Social Media Content Generator (Python)", level=2)
doc.add_paragraph(
    "A Python-based content generator produces branded social media marketing content "
    "for GiveWiZe. It pulls charity data from the Supabase database, enriches it with "
    "live web data, and generates platform-specific content using Claude."
)
doc.add_paragraph("Data sources:")
add_bullets([
    "Supabase database (charity profiles, scores, categories)",
    "Live website scraping (active campaigns, fundraisers, events via keyword detection)",
    "Awareness calendar (40+ cause awareness events mapped to 18 categories by month)",
])
doc.add_paragraph("Five content types:")
add_bullets([
    "Charity Spotlights - feature a specific charity with web-enriched data",
    "Cause Awareness - tied to awareness months/days (e.g., Rare Disease Day, Black History Month)",
    "Campaign Highlights - active fundraisers/events detected from charity websites",
    "Platform Promos - drive traffic to GiveWiZe features (quiz, scores, explore page)",
    "Impact Stats - aggregate data across all charities",
])
doc.add_paragraph(
    "Content is tailored per platform: Instagram posts use emoji, line breaks, 3-5 hashtags, "
    "and a casual tone under 2200 characters. LinkedIn posts are professional, 150-300 words, "
    "with 2-3 hashtags and a data-driven approach. Each post includes a DALL-E image generation prompt."
)
doc.add_paragraph(
    "The generator also detects category gaps based on current awareness events and "
    "auto-onboards charities from ProPublica to fill those gaps, ensuring content never "
    "references a charity not on the platform."
)
doc.add_paragraph(
    "Auto-posting is implemented via LinkedIn OAuth2 API integration. The tool supports "
    "CLI flags for customization: --days, --category, --charity, --post, --post-li, "
    "--no-web, --no-onboard, --no-csv."
)

doc.add_heading("2.8 Frontend Implementation", level=2)
doc.add_paragraph(
    "The frontend consists of 18 pages lazy-loaded via React.lazy() for code splitting, "
    "79 components (37 custom + 42 shadcn/ui with Radix primitives), and 8 custom hooks. "
    "Server state is managed with @tanstack/react-query using a 5-minute staleTime cache "
    "with prefetching on hover for charity detail pages."
)
doc.add_paragraph(
    "The UI uses a dark glassmorphic design system with backdrop-filter blur, semi-transparent "
    "white panels, and gradient backgrounds flowing between sections. The homepage progresses "
    "through Hero, Categories, Featured Charities, How It Works, Testimonials, Call to Action, "
    "and Footer sections."
)
doc.add_paragraph(
    "The optimized bundle sizes after code splitting are approximately 320KB core, 165KB "
    "vendor, 38KB query, and 359KB charts."
)

doc.add_heading("2.9 Email System", level=2)
doc.add_paragraph(
    "The platform uses Resend API for transactional emails with DNS configured (DKIM, SPF, "
    "MX records) on the givewize.org domain. Three email types are implemented:"
)
add_bullets([
    "Charity data request - 30-day tokenized link for charities to submit missing information",
    "Requester notification - branded email when a requested charity is auto-approved with score and profile link",
    "Admin flag notification - alert to info@givewize.org when the pipeline flags a charity for manual review",
])

doc.add_heading("2.10 Authentication & User Profiles", level=2)
doc.add_paragraph(
    "Supabase Auth handles email/password authentication with an AuthContext provider "
    "pattern and protected routes. User profiles support avatar upload to Supabase Storage, "
    "300-character bios, saved favorites, and quiz results. Profile creation is automated "
    "via a Postgres trigger with SECURITY DEFINER."
)

doc.add_heading("2.11 Tool Selection Justifications", level=2)
add_bullets([
    "React/Vite over Next.js: Single-page application was sufficient; no server-side rendering needed; faster development builds",
    "Supabase over Firebase: PostgreSQL for relational queries; built-in auth, edge functions, and Row Level Security",
    "TypeScript: Type safety across 79 components caught bugs early in development",
    "Claude Sonnet over GPT-4: Better structured JSON extraction, lower cost, works within 30-second Edge Function timeout",
    "Resend over SendGrid: Simpler API, straightforward custom domain verification, generous free tier",
    "Python for marketing tool: Rich data science ecosystem, clean scripting syntax, easy API integrations",
])

doc.add_page_break()

# ── 3. Results & Analysis ────────────────────────────────────────────

doc.add_heading("3. Results & Analysis", level=1)

doc.add_heading("3.1 Platform Metrics", level=2)
add_bullets([
    "76 verified charities across 18 cause categories",
    "18 pages, 79 components, 8 custom hooks",
    "GiveWiZe Score range: 2.1 to 4.8 (demonstrates meaningful differentiation)",
    "Onboarding pipeline: ~15-20 seconds per charity, $0.05 cost per charity",
    "Bundle size: ~320KB core (reduced from ~900KB before code splitting)",
    "3 AI agents with non-AI fallbacks achieving 100% pipeline completion rate",
])

doc.add_heading("3.2 Social Media Generator Results", level=2)
add_bullets([
    "Generates 7-day content calendars in approximately 3-5 minutes",
    "Automatically detects awareness events (tested February: American Heart Month, Black History Month, Rare Disease Day, Random Acts of Kindness Week)",
    "Detected active campaigns on 7 charity websites during test runs",
    "Platform-specific content verified with distinct Instagram vs LinkedIn formatting",
    "Successfully auto-posted to LinkedIn via API",
    "CSV export for content scheduling and record-keeping",
])

doc.add_heading("3.3 AI Agent Performance", level=2)
add_bullets([
    "Agent 1 (Enrichment): Extracts 6-8 structured fields from raw HTML in 4-5 seconds with regex fallback",
    "Agent 2 (Quality Review): Correctly flags unverified and low-quality profiles; has overridden incorrect categories",
    "Agent 3 (Email): Generates personalized outreach referencing each charity's actual mission and programs",
    "Content Generator: Produces distinct, non-repetitive posts across all 5 content types with accurate charity data",
])

doc.add_heading("3.4 Scoring Analysis", level=2)
doc.add_paragraph(
    "GiveWiZe Scores show meaningful spread from 2.1 to 4.8, demonstrating that the scoring "
    "methodology effectively differentiates between charities. The null-exclusion approach "
    "prevents unfair penalization of charities with incomplete data, while the longevity and "
    "impact floors at 3.0 ensure newer organizations are not disproportionately ranked low. "
    "The combined score formula (60% GiveWiZe + 40% community) balances internal metrics "
    "with public perception."
)

doc.add_heading("3.5 Limitations", level=2)
add_bullets([
    "Community ratings: 50 of 71 charities share identical scores due to limited source data",
    "Scores computed client-side only, recalculated on each page load",
    "Web scraping is fragile; Charity Navigator and charity website structures can change without notice",
    "Instagram auto-posting blocked: Facebook requires an established account before allowing Page creation",
    "LinkedIn posting currently to personal profile; Company Page posting requires Community Management API approval",
    "DALL-E image generation requires paid OpenAI credits",
    "US-only charity coverage (ProPublica covers US 501(c)(3) organizations only)",
])

doc.add_heading("3.6 What I Would Do Differently", level=2)
add_bullets([
    "Start with independent infrastructure from day one to avoid platform lock-in",
    "Store computed scores in the database to reduce client-side computation",
    "Build review collection earlier to avoid placeholder community ratings",
    "Consider server-side rendering for improved SEO performance",
])

doc.add_page_break()

# ── 4. Discussion ────────────────────────────────────────────────────

doc.add_heading("4. Discussion", level=1)

doc.add_heading("4.1 Technical Challenges and Solutions", level=2)
add_bullets([
    "Lovable Platform Lock-in: The project originally started on Lovable.ai, which provided rapid prototyping but created dependency on their infrastructure. I migrated to an independent Supabase project, learning the importance of owning your infrastructure from the start.",
    "UUID/ID Mismatch: Charity detail pages were completely broken because the database used integer IDs while foreign key references expected UUIDs. This required tracing the data flow through the entire application.",
    "RLS Silent Failures: Supabase Row Level Security queries returned empty arrays instead of errors when permissions were misconfigured, making debugging extremely difficult. I learned to test with the service role key to isolate RLS issues.",
    "Missing Data Devastating Scores: 59 of 71 charities were missing year_founded data, which initially dragged all scores down. The solution was null-exclusion with weight renormalization, which became a core design principle.",
    "Broken Charity Logos: Required three rounds of SQL migration fixes before achieving 68 of 71 working logo URLs.",
    "Edge Function Timeouts: The Python auto-onboarding feature hit 10-second default timeouts when calling Edge Functions. Switching from the Supabase SDK to direct HTTP with 60-second timeouts resolved this.",
    "LinkedIn OAuth: Redirect URI encoding mismatches and scope approval requirements caused repeated failures before successful authentication.",
    "Claude JSON Parsing: The AI model frequently wraps JSON responses in markdown fences or adds preamble text. A 3-strategy extraction approach (direct parse, fence regex, brace matching) handles all cases.",
])

doc.add_heading("4.2 AI Integration Lessons", level=2)
add_bullets([
    "Always build non-AI fallbacks for every AI-powered step (regex extraction, template emails, default logic)",
    "Truncate input to stay within token limits (40KB cap prevents context window overflow)",
    "Parse AI output defensively and never assume clean JSON will be returned",
    "Monitor costs carefully ($0.05 per charity is sustainable at approximately 100 charities per $5 of API credits)",
    'AI-generated content requires explicit guardrails in system prompts (e.g., "do not fabricate statistics")',
])

doc.add_heading("4.3 Digital Marketing Integration", level=2)
doc.add_paragraph(
    "The AI content generator bridges the gap between raw charity data and engaging "
    "donor-facing content. The awareness calendar makes posts timely and culturally "
    "relevant rather than generic. Live website scraping for active campaigns adds urgency "
    "and authenticity. The auto-onboarding feature ensures the platform always has charities "
    "available for any awareness event being promoted. Platform-specific formatting demonstrates "
    "understanding of distinct audience expectations on Instagram versus LinkedIn."
)

doc.add_page_break()

# ── 5. Conclusion ────────────────────────────────────────────────────

doc.add_heading("5. Conclusion", level=1)
doc.add_paragraph(
    "GiveWiZe demonstrates how AI can be integrated throughout a full-stack application "
    "to solve real problems in the charitable giving space. The platform combines a "
    "React/TypeScript frontend, Supabase backend, three Claude AI agents for automated "
    "charity onboarding, and a Python-based social media content generator with LinkedIn "
    "API integration."
)
doc.add_paragraph(
    "The key technical achievement is end-to-end automation: from a user submitting a "
    "charity request to the charity being verified, scored, profiled, quality-reviewed, "
    "and added to the platform in under 20 seconds, followed by AI-generated marketing "
    "content that can be auto-posted to social media."
)
doc.add_paragraph(
    "This project provided significant technical growth across React, TypeScript, Supabase, "
    "AI integration, API development, Python scripting, OAuth flows, web scraping, and email "
    "systems. The platform is live with real charity data, real scoring, and posts going to "
    "real social media accounts."
)
doc.add_paragraph(
    "Future work includes Instagram integration once the Facebook Page requirement is met, "
    "LinkedIn Company Page posting pending API approval, a mobile application, donation "
    "tracking integration, and expanded international charity coverage."
)

doc.add_page_break()

# ── Visual Elements ──────────────────────────────────────────────────

doc.add_heading("Visual Elements to Include", level=1)
add_bullets([
    "[INSERT: System architecture diagram showing all components and data flow]",
    "[INSERT: Onboarding pipeline flowchart with AI agent callouts]",
    "[INSERT: Scoring formula diagram with weights and sub-scores]",
    "[INSERT: Screenshot - Homepage hero section]",
    "[INSERT: Screenshot - Charity detail page with scores]",
    "[INSERT: Screenshot - Quiz flow (question and results)]",
    "[INSERT: Screenshot - Request a Charity page with real-time feedback]",
    "[INSERT: Screenshot - Terminal showing social media generator running]",
    "[INSERT: Screenshot - LinkedIn post that was auto-posted]",
    "[INSERT: Screenshot - Admin flag email received in inbox]",
    "[INSERT: Screenshot - Supabase dashboard]",
    "[INSERT: Code snippet - Scoring computation]",
    "[INSERT: Code snippet - AI agent prompt]",
    "[INSERT: Code snippet - Social media generator Claude call]",
    "[INSERT: Comparison table - GiveWiZe vs existing platforms]",
])

# ── Save ─────────────────────────────────────────────────────────────

filepath = r"C:\Users\hanna\givewize-charity-explorer\docs\GiveWiZe_Technical_Report_Draft.docx"
doc.save(filepath)
print(f"Saved to {filepath}")
